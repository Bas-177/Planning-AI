"""
FastAPI applicatie voor planningsmodule
"""
from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from starlette.templating import Jinja2Templates
from starlette.requests import Request
from typing import List, Optional
from datetime import datetime, date
import re
import io

from app.models import Order, OrderUpdate, Standard, Suggestion, VrijeDag, Medewerker, WeekPlanning
from app.database import ExcelDatabase
from app.ai_suggestions import AISuggestionEngine

app = FastAPI(title="Planning Industrie AI", version="1.0.0")

# Static files en templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# Database
db = ExcelDatabase()
ai_engine = AISuggestionEngine(db)


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Hoofdpagina"""
    orders = db.get_all_orders()
    stats = {
        'totaal_orders': len(orders),
        'open_orders': len([o for o in orders if not o.get('project_gereed', False)]),
        'deadline_nadert': len([o for o in orders if _deadline_nadert(o)]),
    }
    return templates.TemplateResponse("index.html", {"request": request, "stats": stats})


@app.get("/orders", response_class=HTMLResponse)
async def orders_page(request: Request):
    """Orders pagina"""
    return templates.TemplateResponse("orders.html", {"request": request})


@app.get("/planning", response_class=HTMLResponse)
async def planning_page(request: Request):
    """Planning pagina (Kanban/Tabel view)"""
    return templates.TemplateResponse("planning.html", {"request": request})


@app.get("/planning/week", response_class=HTMLResponse)
async def planning_week_page(request: Request):
    """Planning pagina (Week view)"""
    return templates.TemplateResponse("planning_week.html", {"request": request})


@app.get("/api/orders", response_model=List[dict])
async def get_orders():
    """Haal alle orders op"""
    try:
        orders = db.get_all_orders()
        if orders is None:
            orders = []
        # Converteer datetime objecten naar strings
        for order in orders:
            for key, value in list(order.items()):
                if isinstance(value, (datetime, date)):
                    order[key] = value.isoformat() if hasattr(value, 'isoformat') else str(value)
        return orders
    except Exception as e:
        import logging
        logging.error(f"Fout bij ophalen orders: {e}")
        return []


@app.get("/api/orders/{ordernummer}", response_model=dict)
async def get_order(ordernummer: str):
    """Haal specifieke order op"""
    order = db.get_order(ordernummer)
    if not order:
        raise HTTPException(status_code=404, detail="Order niet gevonden")
    return order


@app.post("/api/orders", response_model=dict)
async def create_order(order: Order):
    """Maak nieuwe order aan"""
    try:
        result = db.create_order(order)
        # Converteer datetime/date objecten naar strings voor JSON serialisatie
        from datetime import datetime, date
        cleaned_result = {}
        for key, value in result.items():
            if value is None:
                cleaned_result[key] = None
            elif isinstance(value, (datetime, date)):
                cleaned_result[key] = value.isoformat() if hasattr(value, 'isoformat') else str(value)
            elif isinstance(value, float) and (value != value or value == float('inf') or value == float('-inf')):
                cleaned_result[key] = None if key in ['start_datum', 'eind_datum'] else 0.0
            else:
                cleaned_result[key] = value
        return cleaned_result
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        import logging
        logging.error(f"Fout bij aanmaken order: {e}")
        raise HTTPException(status_code=500, detail=f"Fout bij aanmaken order: {str(e)}")


@app.put("/api/orders/{ordernummer}", response_model=dict)
async def update_order(ordernummer: str, order_update: dict):
    """Update bestaande order"""
    try:
        # Accepteer dict direct (vanuit JavaScript)
        updated = db.update_order(ordernummer, order_update)
        if not updated:
            raise HTTPException(status_code=404, detail="Order niet gevonden")
        
        # Converteer datetime objecten naar strings
        for key, value in list(updated.items()):
            if isinstance(value, (datetime, date)):
                updated[key] = value.isoformat() if hasattr(value, 'isoformat') else str(value)
        
        return updated
    except HTTPException:
        raise
    except Exception as e:
        import logging
        logging.error(f"Fout bij updaten order {ordernummer}: {e}")
        raise HTTPException(status_code=500, detail=f"Fout bij updaten order: {str(e)}")


@app.delete("/api/orders/{ordernummer}")
async def delete_order(ordernummer: str):
    """Verwijder order"""
    success = db.delete_order(ordernummer)
    if not success:
        raise HTTPException(status_code=404, detail="Order niet gevonden")
    return {"message": "Order verwijderd"}


@app.get("/api/search")
async def search_orders(q: str = Query(..., description="Zoekterm")):
    """Zoek in orders"""
    return db.search_orders(q)


@app.get("/api/suggestions/{ordernummer}", response_model=List[Suggestion])
async def get_suggestions(ordernummer: str):
    """Haal AI suggesties op voor order"""
    order = db.get_order(ordernummer)
    if not order:
        raise HTTPException(status_code=404, detail="Order niet gevonden")
    
    return ai_engine.get_suggestions(order)


@app.get("/api/planning")
async def get_planning():
    """Haal planningsoverzicht op (Kanban view)"""
    try:
        orders = db.get_all_orders()
        if orders is None:
            orders = []
        
        # Groepeer per status - LOGICA AANGEPAST
        # Materiaal besteld = besteld (niet binnen)
        # Materiaal binnen = binnen (niet in productie)
        planning = {
            'werkvoorbereiding': [],
            'materiaal_besteld': [],
            'materiaal_binnen': [],
            'in_productie': [],
            'productie_gereed': [],
            'project_gereed': []
        }
        
        for order in orders:
            try:
                # Converteer datetime/date objecten naar strings voor JSON serialisatie
                order_clean = {}
                for key, value in order.items():
                    if value is None:
                        order_clean[key] = None
                    elif isinstance(value, (datetime, date)):
                        order_clean[key] = value.isoformat() if hasattr(value, 'isoformat') else str(value)
                    else:
                        order_clean[key] = value
                
                # Parse boolean values die als string kunnen komen
                project_gereed = order_clean.get('project_gereed', False)
                if isinstance(project_gereed, str):
                    project_gereed = project_gereed.lower() in ['true', '1', 'yes']
                elif project_gereed is None:
                    project_gereed = False
                
                productie_gereed = order_clean.get('productie_gereed', False)
                if isinstance(productie_gereed, str):
                    productie_gereed = productie_gereed.lower() in ['true', '1', 'yes']
                elif productie_gereed is None:
                    productie_gereed = False
                
                materiaal_binnen = order_clean.get('materiaal_binnen', False)
                if isinstance(materiaal_binnen, str):
                    materiaal_binnen = materiaal_binnen.lower() in ['true', '1', 'yes']
                elif materiaal_binnen is None:
                    materiaal_binnen = False
                
                materiaal_besteld = order_clean.get('materiaal_besteld', False)
                if isinstance(materiaal_besteld, str):
                    materiaal_besteld = materiaal_besteld.lower() in ['true', '1', 'yes']
                elif materiaal_besteld is None:
                    materiaal_besteld = False
                
                if project_gereed:
                    planning['project_gereed'].append(order_clean)
                elif productie_gereed:
                    planning['productie_gereed'].append(order_clean)
                elif materiaal_binnen:
                    planning['in_productie'].append(order_clean)
                elif materiaal_besteld:
                    planning['materiaal_besteld'].append(order_clean)
                else:
                    planning['werkvoorbereiding'].append(order_clean)
            except Exception as order_error:
                # Skip deze order als er een fout is, maar log het
                import logging
                logging.error(f"Fout bij verwerken order {order.get('ordernummer', 'unknown')}: {order_error}")
                continue
        
        return planning
    except Exception as e:
        import traceback
        import logging
        error_detail = f"Fout bij laden planning: {str(e)}"
        logging.error(f"{error_detail}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=error_detail)


@app.get("/api/planning/week")
async def get_week_planning(week_nummer: Optional[int] = None, jaar: Optional[int] = None):
    """Haal week planning op met medewerkers en toewijzingen"""
    from datetime import datetime
    
    try:
        # Bepaal week en jaar
        if not week_nummer or not jaar:
            now = datetime.now()
            jaar = jaar or now.year
            week_nummer = week_nummer or now.isocalendar()[1]
        
        # Haal alle data op
        orders = db.get_all_orders()
        medewerkers = db.get_medewerkers()
        assignments = db.get_order_assignments()
        week_planning = db.get_week_planning(week_nummer, jaar)
        
        # Maak planning structuur
        planning_data = {
            'week_nummer': week_nummer,
            'jaar': jaar,
            'medewerkers': []
        }
        
        # Voor elke actieve medewerker
        for medewerker in [m for m in medewerkers if m.get('actief')]:
            medewerker_planning = {
                'naam': medewerker['naam'],
                'beschikbare_uren': next((wp['beschikbare_uren'] for wp in week_planning if wp['medewerker'] == medewerker['naam']), medewerker.get('standaard_uren_per_week', 40)),
                'toewijzingen': []
            }
            
            # Zoek toewijzingen voor deze medewerker
            for assignment in assignments:
                if assignment['medewerker'] == medewerker['naam']:
                    # Check of toewijzing in deze week valt
                    if assignment.get('week_nummer') == week_nummer and assignment.get('jaar') == jaar:
                        # Zoek order details
                        order = next((o for o in orders if o['ordernummer'] == assignment['ordernummer']), None)
                        if order:
                            medewerker_planning['toewijzingen'].append({
                                'ordernummer': assignment['ordernummer'],
                                'bewerking': assignment['bewerking'],
                                'uren': assignment['uren'],
                                'order_omschrijving': order.get('omschrijving', ''),
                                'klant': order.get('klant', '')
                            })
            
            planning_data['medewerkers'].append(medewerker_planning)
        
        return planning_data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Fout bij laden week planning: {str(e)}")


@app.get("/api/order-assignments")
async def get_order_assignments(ordernummer: Optional[str] = None, medewerker: Optional[str] = None):
    """Haal order toewijzingen op"""
    try:
        assignments = db.get_order_assignments(ordernummer, medewerker)
        # Zorg dat alle waarden JSON-serialiseerbaar zijn
        cleaned = []
        for assignment in assignments:
            cleaned_assignment = {}
            for key, value in assignment.items():
                if value is None:
                    cleaned_assignment[key] = None
                elif isinstance(value, (datetime, date)):
                    cleaned_assignment[key] = value.isoformat() if hasattr(value, 'isoformat') else str(value)
                elif isinstance(value, float) and (value != value or value == float('inf') or value == float('-inf')):
                    # NaN of infinity
                    cleaned_assignment[key] = None if key in ['start_datum', 'eind_datum'] else 0.0
                else:
                    cleaned_assignment[key] = value
            cleaned.append(cleaned_assignment)
        return cleaned
    except Exception as e:
        import logging
        logging.error(f"Fout bij ophalen order assignments: {e}")
        raise HTTPException(status_code=500, detail=f"Fout bij ophalen data: {str(e)}")


@app.post("/api/order-assignments")
async def create_order_assignment(assignment: dict):
    """Voeg order toewijzing toe"""
    return db.add_order_assignment(assignment)


@app.delete("/api/order-assignments/{ordernummer}/{medewerker}/{bewerking}")
async def delete_order_assignment(ordernummer: str, medewerker: str, bewerking: str):
    """Verwijder order toewijzing"""
    success = db.delete_order_assignment(ordernummer, medewerker, bewerking)
    if not success:
        raise HTTPException(status_code=404, detail="Toewijzing niet gevonden")
    return {"message": "Toewijzing verwijderd"}


@app.get("/api/standards", response_model=List[dict])
async def get_standards():
    """Haal standaard doorlooptijden op"""
    return db.get_standards()


@app.post("/api/standards", response_model=dict)
async def create_standard(standard: Standard):
    """Voeg standaard doorlooptijd toe"""
    return db.create_standard(standard)


@app.delete("/api/standards/{nabehandeling}")
async def delete_standard(nabehandeling: str):
    """Verwijder standaard doorlooptijd"""
    success = db.delete_standard(nabehandeling)
    if not success:
        raise HTTPException(status_code=404, detail="Standaard niet gevonden")
    return {"message": "Standaard verwijderd"}


@app.get("/api/notifications")
async def get_notifications():
    """Haal meldingen op"""
    orders = db.get_all_orders()
    notifications = []
    
    for order in orders:
        # Deadline nadert
        if _deadline_nadert(order):
            notifications.append({
                'type': 'deadline',
                'prioriteit': 'high',
                'bericht': f"Order {order['ordernummer']} heeft deadline binnen 3 dagen",
                'ordernummer': order['ordernummer']
            })
        
        # Materiaal nog niet besteld maar productie start binnenkort
        if not order.get('materiaal_besteld') and _productie_start_binnenkort(order):
            notifications.append({
                'type': 'materiaal',
                'prioriteit': 'medium',
                'bericht': f"Order {order['ordernummer']}: Materiaal nog niet besteld",
                'ordernummer': order['ordernummer']
            })
    
    return notifications


def _deadline_nadert(order: dict) -> bool:
    """Check of deadline binnen 3 dagen is"""
    leverdatum = order.get('leverdatum')
    if not leverdatum:
        return False
    
    if isinstance(leverdatum, str):
        leverdatum = datetime.fromisoformat(leverdatum).date()
    elif isinstance(leverdatum, datetime):
        leverdatum = leverdatum.date()
    
    dagen_tot_deadline = (leverdatum - datetime.now().date()).days
    return 0 <= dagen_tot_deadline <= 3


def _productie_start_binnenkort(order: dict) -> bool:
    """Check of productie binnen 5 dagen start"""
    leverdatum = order.get('leverdatum')
    if not leverdatum:
        return False
    
    if isinstance(leverdatum, str):
        leverdatum = datetime.fromisoformat(leverdatum).date()
    elif isinstance(leverdatum, datetime):
        leverdatum = leverdatum.date()
    
    dagen_tot_deadline = (leverdatum - datetime.now().date()).days
    return 0 <= dagen_tot_deadline <= 5


# Vrije dagen endpoints
@app.get("/api/vrije-dagen", response_model=List[dict])
async def get_vrije_dagen():
    """Haal alle vrije dagen op"""
    return db.get_vrije_dagen()


@app.post("/api/vrije-dagen", response_model=dict)
async def create_vrije_dag(vrije_dag: VrijeDag):
    """Voeg vrije dag toe"""
    vrije_dag_dict = vrije_dag.model_dump()
    return db.add_vrije_dag(vrije_dag_dict)


@app.delete("/api/vrije-dagen/{datum}")
async def delete_vrije_dag(datum: str):
    """Verwijder vrije dag"""
    success = db.delete_vrije_dag(datum)
    if not success:
        raise HTTPException(status_code=404, detail="Vrije dag niet gevonden")
    return {"message": "Vrije dag verwijderd"}


# Medewerkers endpoints
@app.get("/api/medewerkers", response_model=List[dict])
async def get_medewerkers():
    """Haal alle medewerkers op"""
    return db.get_medewerkers()


@app.post("/api/medewerkers", response_model=dict)
async def create_medewerker(medewerker: Medewerker):
    """Voeg medewerker toe"""
    try:
        medewerker_dict = medewerker.model_dump()
        return db.add_medewerker(medewerker_dict)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/api/medewerkers/{naam}", response_model=dict)
async def update_medewerker(naam: str, update_data: dict):
    """Update medewerker"""
    updated = db.update_medewerker(naam, update_data)
    if not updated:
        raise HTTPException(status_code=404, detail="Medewerker niet gevonden")
    return updated


@app.delete("/api/medewerkers/{naam}")
async def delete_medewerker(naam: str):
    """Verwijder medewerker"""
    success = db.delete_medewerker(naam)
    if not success:
        raise HTTPException(status_code=404, detail="Medewerker niet gevonden")
    return {"message": "Medewerker verwijderd"}


# Week planning endpoints
@app.get("/api/week-planning")
async def get_week_planning(week_nummer: Optional[int] = None, jaar: Optional[int] = None):
    """Haal week planning op"""
    return db.get_week_planning(week_nummer, jaar)


@app.post("/api/week-planning", response_model=dict)
async def set_week_planning(week_planning: WeekPlanning):
    """Stel week planning in"""
    week_planning_dict = week_planning.model_dump()
    return db.set_week_planning(week_planning_dict)


@app.get("/agenda", response_class=HTMLResponse)
async def agenda_page(request: Request):
    """Agenda pagina"""
    return templates.TemplateResponse("agenda.html", {"request": request})


@app.get("/medewerkers", response_class=HTMLResponse)
async def medewerkers_page(request: Request):
    """Medewerkers planning pagina"""
    return templates.TemplateResponse("medewerkers.html", {"request": request})


@app.get("/data", response_class=HTMLResponse)
async def data_page(request: Request):
    """Data pagina voor standaard doorlooptijden"""
    return templates.TemplateResponse("data.html", {"request": request})


@app.get("/projectplanning", response_class=HTMLResponse)
async def projectplanning_page(request: Request):
    """Projectplanning pagina - alle projecten onder elkaar"""
    return templates.TemplateResponse("projectplanning.html", {"request": request})


# OCR endpoint - vereist python-multipart
# Check of python-multipart beschikbaar is
try:
    import multipart  # Check of python-multipart geïnstalleerd is
    from fastapi import UploadFile, File
    
    @app.post("/api/ocr/screenshot")
    async def ocr_screenshot(file: UploadFile = File(...)):
        """Verwerk screenshot/PDF/Word bestand - accepteert 'file' parameter"""
        """Verwerk screenshot/PDF/Word bestand en probeer gegevens uit te lezen"""
        try:
            # Lees bestand data
            file_data = await file.read()
            content_type = file.content_type or ''
            file_name = file.filename or ''
            
            result = {
                "projectnummer": None,
                "omschrijving": None,
                "klant": None,
                "leverdatum": None,
                "startdatum": None,
                "referentienummer": None,
                "message": None
            }
            
            text = ""
            
            # Handle Word documents (.docx en .doc)
            # Check zowel MIME type als bestandsnaam extensie
            is_word_doc = (
                content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
                content_type == 'application/msword' or
                file_name.lower().endswith('.docx') or
                file_name.lower().endswith('.doc')
            )
            
            if is_word_doc:
                try:
                    from docx import Document
                    from io import BytesIO
                    doc = Document(BytesIO(file_data))
                    text = '\n'.join([para.text for para in doc.paragraphs])
                    if not text or len(text.strip()) < 3:
                        # Probeer ook tabellen
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = '\n'.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                                if row_text:
                                    text += '\n' + row_text
                    result["message"] = f"Word document gelezen: {len(text)} karakters"
                    result["fileType"] = "Word document"  # Voor frontend feedback
                except ImportError:
                    import logging
                    logging.warning("python-docx niet geïnstalleerd. Installeer met: pip install python-docx")
                    result["message"] = "python-docx niet geïnstalleerd. Installeer met: pip install python-docx"
                    # Als Word bestand wordt geüpload maar python-docx niet geïnstalleerd is, return error
                    return result
                except Exception as e:
                    import logging
                    logging.error(f"Fout bij lezen Word document: {e}", exc_info=True)
                    result["message"] = f"Fout bij lezen Word document: {str(e)}. Zorg dat het bestand een geldig Word document is."
                    return result
            
            # Handle PDF files
            elif content_type == 'application/pdf' or file_name.lower().endswith('.pdf'):
                try:
                    import PyPDF2
                    from io import BytesIO
                    pdf = PyPDF2.PdfReader(BytesIO(file_data))
                    text = '\n'.join([page.extract_text() for page in pdf.pages])
                    result["message"] = f"PDF gelezen: {len(text)} karakters"
                except ImportError:
                    result["message"] = "PyPDF2 niet geïnstalleerd. Installeer met: pip install PyPDF2"
                    return result
                except Exception as e:
                    result["message"] = f"Fout bij lezen PDF: {str(e)}"
                    return result
            
            # Handle images (existing logic)
            elif content_type.startswith('image/'):
                image_data = file_data
                
                # Probeer OCR met pytesseract (als beschikbaar)
                try:
                    from PIL import Image
                    import io
                    import re
                    import logging
                    import subprocess
                    import os
                    
                    # Open image
                    img = Image.open(io.BytesIO(image_data))
                    
                    # Probeer pytesseract (als geïnstalleerd)
                    try:
                        import pytesseract
                        
                        # Configureer Tesseract pad (probeer eerst PATH, dan standaard pad)
                        tesseract_found = False
                        try:
                            subprocess.run(['tesseract', '--version'], 
                                         capture_output=True, 
                                         timeout=5,
                                         check=True)
                            tesseract_found = True
                            logging.info("Tesseract gevonden in PATH")
                        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                            default_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
                            if os.path.exists(default_path):
                                pytesseract.pytesseract.tesseract_cmd = default_path
                                tesseract_found = True
                                logging.info(f"Tesseract gevonden op: {default_path}")
                            else:
                                alt_paths = [
                                    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                                    r'C:\Tesseract-OCR\tesseract.exe',
                                ]
                                for alt_path in alt_paths:
                                    if os.path.exists(alt_path):
                                        pytesseract.pytesseract.tesseract_cmd = alt_path
                                        tesseract_found = True
                                        logging.info(f"Tesseract gevonden op: {alt_path}")
                                        break
                        
                        if not tesseract_found:
                            raise FileNotFoundError(
                                "Tesseract niet gevonden. Installeer Tesseract of configureer het pad handmatig."
                            )

                        # Extract tekst uit afbeelding
                        logging.info(f"Start OCR extractie voor afbeelding {img.size}")
                        text = pytesseract.image_to_string(img, lang='nld+eng')
                        logging.info(f"OCR tekst geëxtraheerd: {len(text)} karakters")
                        
                    except ImportError:
                        logging.warning("pytesseract niet geïnstalleerd")
                        text = ""
                    except Exception as e:
                        logging.error(f"OCR fout: {e}")
                        text = ""
                except Exception as e:
                    logging.error(f"Fout bij image processing: {e}")
                    text = ""
            else:
                result["message"] = "Bestandstype niet ondersteund. Gebruik afbeelding, PDF of Word document."
                return result
            
            # Parse tekst voor order informatie (als tekst beschikbaar is)
            if text and len(text.strip()) >= 3:
                # Voer altijd extractie uit als er tekst is
                parsed_result = extract_order_info_from_text(text)
                result.update(parsed_result)  # Update result met parsed data
                
                # Update message alleen als er nog geen message is (bijv. van Word lezen)
                if not result.get("message"):
                    if any([result.get("projectnummer"), result.get("omschrijving"), result.get("klant"), result.get("leverdatum")]):
                        result["message"] = f"Document succesvol verwerkt: {len(text)} karakters gelezen"
                    else:
                        result["message"] = f"Geen ordergegevens gevonden in document. {len(text)} karakters gelezen maar kon geen projectnummer, klant of datums vinden. Bekijk de tekst in de console voor debugging."
                else:
                    # Als er al een message is (bijv. "Word document gelezen"), voeg extractie resultaat toe
                    extracted_fields = []
                    if result.get("projectnummer"):
                        extracted_fields.append(f"Projectnummer: {result['projectnummer']}")
                    if result.get("omschrijving"):
                        extracted_fields.append(f"Omschrijving: {result['omschrijving']}")
                    if result.get("klant"):
                        extracted_fields.append(f"Klant: {result['klant']}")
                    if result.get("leverdatum"):
                        extracted_fields.append(f"Leverdatum: {result['leverdatum']}")
                    if extracted_fields:
                        result["message"] = result["message"] + " | " + " | ".join(extracted_fields)
            elif not text:
                if not result.get("message"):
                    result["message"] = "Geen tekst kunnen extraheren uit document. Zorg dat het document leesbaar is."
            
            return result
        except Exception as e:
            import logging
            logging.error(f"Fout bij OCR verwerking: {e}")
            raise HTTPException(status_code=500, detail=f"Fout bij verwerken screenshot: {str(e)}")
except ImportError:
    # OCR endpoint niet beschikbaar als python-multipart niet geïnstalleerd is
    @app.post("/api/ocr/screenshot")
    async def ocr_screenshot():
        """OCR functionaliteit niet beschikbaar - installeer python-multipart"""
        raise HTTPException(
            status_code=503, 
            detail="OCR functionaliteit vereist python-multipart. Installeer met: pip install python-multipart"
        )


# Helper functie om order info uit tekst te extraheren
def extract_order_info_from_text(text: str) -> dict:
    """
    Extract order informatie uit OCR tekst
    Zoekt naar specifieke labels en extraheert de waarden ernaast (onafhankelijk van positie):
    - "Project" / "Ons referentie nr." → Ordernummer
    - "Omschrijving" / "Betreft" → Korte Omschrijving
    - "Relatie" → Klant
    - "Referentienr." / "Uw referentie" → Klantreferentie
    - "Opleverdatum" / "Leverdatum" → Leverdatum
    """
    result = {
        "projectnummer": None,
        "omschrijving": None,
        "klant": None,
        "leverdatum": None,
        "startdatum": None,
        "referentienummer": None,
        "message": None
    }
    
    if not text:
        return result
    
    text_upper = text.upper()
    text_lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # ========== 1. PROJECT = ORDERNUMMER ==========
    # Zoek naar label "Project" en haal waarde ernaast op (P##-#### formaat)
    # Patronen: "Project: P25-0061", "Project P25-0061", "Offertenr.: P25-0061"
    
    project_label_patterns = [
        r'Project[:\s]+(P[0-9OIl]{2}-[0-9OIl]{4,})',  # Project: P25-0061 (HOOGSTE PRIORITEIT)
        r'Offertenr\.?[:\s]+(P[0-9OIl]{2}-[0-9OIl]{4,})',  # Offertenr.: P25-0061
        r'Projectnr\.?[:\s]+(P[0-9OIl]{2}-[0-9OIl]{4,})',  # Projectnr.: P25-0061
    ]
    
    # Zoek eerst naar label "Project" (onafhankelijk van positie)
    for pattern in project_label_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            project_num = match.group(1).strip().replace(' ', '').replace('.', '').replace(',', '')
            # Corrigeer OCR fouten: O -> 0, I -> 1, l -> 1
            project_num = project_num.replace('O', '0').replace('I', '1').replace('l', '1').replace('|', '1')
            # Validatie: moet P gevolgd door 2 cijfers, dan -, dan 4+ cijfers
            if re.match(r'^P\d{2}-\d{4,}$', project_num, re.IGNORECASE):
                result["projectnummer"] = project_num
                break
    
    # Fallback: zoek direct P##-#### patroon (zonder label)
    if not result["projectnummer"]:
        project_patterns = [
            r'\b(P[0-9OIl]{2}-[0-9OIl]{4})\b',  # P25-0061 (exact 4 cijfers)
            r'\b(P\d{2}-\d{4,})\b',  # P25-0061 (4+ cijfers)
        ]
        for pattern in project_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                project_num = match.group(1).strip().replace(' ', '').replace('.', '')
                project_num = project_num.replace('O', '0').replace('I', '1').replace('l', '1').replace('|', '1')
                if re.match(r'^P\d{2}-\d{4,}$', project_num, re.IGNORECASE):
                    result["projectnummer"] = project_num
                    break
    
    # Zoek klant/relatie - BELANGRIJK: zoek specifiek het veld direct na "Relatie:" label
    # In het screenshot staat dit als: "Relatie: Van Doorn Infra" (stopt bij "Contact pers.")
    # Filter alle verkeerde matches zoals "Historie", "Algemene", etc.
    
    # Strategie: zoek per regel naar "Relatie:" en neem de tekst tot aan "Contact"
    for line in text_lines:
        # Zoek "Relatie:" op regel
        relatie_match = re.search(r'Relatie[:\s]+([^\n]{3,80})', line, re.IGNORECASE)
        if relatie_match:
            klant_raw = relatie_match.group(1).strip()
            # Stop bij "Contact pers." of "Contact"
            klant = re.split(r'\s+Contact', klant_raw, flags=re.IGNORECASE)[0].strip()
            klant = re.split(r'\s+pers\.', klant, flags=re.IGNORECASE)[0].strip()
            
            # Schoon op: verwijder extra whitespace
            klant = re.sub(r'\s+', ' ', klant)
            
            # Filter verkeerde matches - deze woorden duiden op andere velden, niet klant
            exclude_patterns = [
                r'^Historie',
                r'^Algemene',
                r'Activiteiten',
                r'Historie.*Algemene',
                r'Algemene.*Historie',
            ]
            is_valid = True
            for exclude_pattern in exclude_patterns:
                if re.search(exclude_pattern, klant, re.IGNORECASE):
                    is_valid = False
                    break
            
            # Validatie: moet beginnen met hoofdletter, redelijke lengte, en geen nummers aan begin
            if is_valid and len(klant) >= 3 and len(klant) <= 80:
                if re.match(r'^[A-Z]', klant) and not re.match(r'^\d', klant):
                    # Verwijder trailing speciale karakters
                    klant = re.sub(r'[.,]+$', '', klant).strip()
                    result["klant"] = klant
                    break
    
    # Fallback: zoek algemene "Relatie:" pattern in hele tekst (als regel-gebaseerd niet werkt)
    if not result["klant"]:
        klant_patterns = [
            # Relatie: [tekst] (stop bij Contact pers. of nieuwe regel) - HOOGSTE PRIORITEIT
            r'Relatie[:\s]+([A-Z][a-zA-Z\s]{2,60}?)(?:\s+Contact|Contact\s+pers|pers\.|$|\n)',
            # Relatie: [tekst] met meerdere woorden
            r'Relatie[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,4})(?:\s+Contact|$|\n)',
        ]
        
        for pattern in klant_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                klant = match.group(1).strip()
                # Filter verkeerde matches
                if not re.search(r'(Historie|Algemene|Activiteiten)', klant, re.IGNORECASE):
                    klant = re.sub(r'\s+', ' ', klant)
                    klant = re.split(r'\s+Contact', klant, flags=re.IGNORECASE)[0].strip()
                    if len(klant) >= 3 and len(klant) <= 80 and re.match(r'^[A-Z]', klant):
                        result["klant"] = klant
                        break
    
    # Fallback: zoek naar bekende klantnamen in tekst
    if not result["klant"]:
        # Zoek patronen zoals "Van Doorn Infra" (woorden met hoofdletters)
        klant_match = re.search(r'\b([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b', text)
        if klant_match:
            klant = klant_match.group(1).strip()
            # Filter verkeerde matches (bijv. "Project type" of "Contact pers")
            if not re.search(r'(Project|Contact|Fase|Status|Email|Telefoon|pers)', klant, re.IGNORECASE):
                if 3 <= len(klant) <= 80:
                    result["klant"] = klant
    
    # Zoek omschrijving - BELANGRIJK: eerste zin/veld, niet te lang
    # In dit systeem staat omschrijving vaak als eerste veld na projectnummer
    # Ook zoeken naar "Betreft:" label (gebruikt in opdrachtbevestigingen)
    omschrijving_patterns = [
        # HOOGSTE PRIORITEIT: "Betreft:" label (gebruikt in opdrachtbevestigingen) - "Betreft: Koker 50x50x3"
        r'Betreft[:\s]+([^\n]{3,80}?)(?:\s*(?:Ons\s+referentie|Ons\s+Referentie|Uw\s+referentie|Orderdatum|Leverdatum|Project\s+type|Project|Relatie|Klant|Fase|Status|Contact|$|\n))',
        # "Omschrijving:" gevolgd door tekst (stop bij "Project type" of volgende label)
        r'(?:Omschrijving|Beschrijving)[:\s]+([^\n]{3,60}?)(?:\s*(?:Project\s+type|Project|Relatie|Klant|Fase|Status|Contact|$|\n))',
        # Zoek tekst direct na projectnummer (eerste zin, stop bij "Project type")
        r'(?:Project|P\d{2}-\d{4,})[^\n]{0,10}[\s:]+([A-Z][^\n]{3,60}?)(?:\s+(?:Project\s+type|Project|Relatie|Klant|Fase|Status|Contact|$|\n))',
    ]
    
    for pattern in omschrijving_patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            omschrijving = match.group(1).strip()
            # Schoon op: verwijder speciale tekens aan begin/einde (maar behoud cijfers en speciale karakters zoals x voor maten)
            omschrijving = re.sub(r'^[^\w0-9xX]+|[^\w0-9xX]+$', '', omschrijving)  # Verwijder niet-woorden/cijfers aan begin/einde
            omschrijving = re.sub(r'\[|\]', '', omschrijving)  # Verwijder brackets
            
            # Voor "Betreft:" - stop bij volgende velden zoals "Ons referentie", "Uw referentie", etc.
            if 'ons referentie' in text.lower() or 'uw referentie' in text.lower():
                # Stop bij "Ons referentie" of "Uw referentie"
                omschrijving = re.split(r'\s+(?:Ons|Uw)\s+referentie', omschrijving, flags=re.IGNORECASE)[0].strip()
            
            # Stop bij "Project type" (hoogste prioriteit) - voorkomt "[Frame van Doorn Project type Regiewerk a"
            # Dit is het belangrijkste stopwoord!
            if 'project' in omschrijving.lower() and 'type' in omschrijving.lower():
                # Split op "Project type" of "Project"
                parts = re.split(r'\s+Project\s+type', omschrijving, flags=re.IGNORECASE)
                if len(parts) > 1 and len(parts[0].strip()) >= 3:
                    omschrijving = parts[0].strip()
                else:
                    # Als "Project type" niet werkt, probeer alleen "Project"
                    parts = re.split(r'\s+Project\s+', omschrijving, flags=re.IGNORECASE)
                    if len(parts) > 1 and len(parts[0].strip()) >= 3:
                        omschrijving = parts[0].strip()
            else:
                # Stop bij andere woorden die nieuwe velden aangeven
                stop_words = ['Regiewerk', 'Relatie', 'Klant', 'Fase', 'Status', 'Contact', 'Orderdatum', 'Leverdatum']
                for stop_word in stop_words:
                    if stop_word.lower() in omschrijving.lower():
                        parts = re.split(rf'\s+{re.escape(stop_word)}', omschrijving, flags=re.IGNORECASE)
                        if parts and len(parts[0].strip()) >= 3:
                            omschrijving = parts[0].strip()
                            break
            
            # Limiteer lengte (max 100 karakters voor omschrijving - iets meer voor technische omschrijvingen)
            if len(omschrijving) > 100:
                # Stop bij laatste volledige woord voor 100 karakters
                omschrijving = omschrijving[:100].rsplit(' ', 1)[0] if ' ' in omschrijving[:100] else omschrijving[:100]
            
            # Verwijder trailing speciale karakters en brackets (maar behoud x voor maten)
            omschrijving = re.sub(r'[^\w\s0-9xX]+$', '', omschrijving).strip()
            
            # Validatie: moet minimaal 3 karakters zijn, maximaal 100
            if len(omschrijving) >= 3 and len(omschrijving) <= 100:
                result["omschrijving"] = omschrijving
                break
    
    # Fallback: zoek eerste zin die lijkt op omschrijving (geen nummer, geen datum)
    if not result["omschrijving"]:
        # Zoek naar tekst die begint met hoofdletter of brackets, 3-50 karakters
        # Bijvoorbeeld: "[Frame van Doorn" of "Frame van Doorn"
        fallback_patterns = [
            r'\[([A-Z][a-z]+(?:\s+[a-z]+){1,5})',  # [Frame van Doorn
            r'([A-Z][a-z]+(?:\s+[a-z]+){1,5})(?:\s+Project|\s+type|$)',  # Frame van Doorn (stop bij Project type)
        ]
        for pattern in fallback_patterns:
            fallback_match = re.search(pattern, text, re.IGNORECASE)
            if fallback_match:
                omschrijving = fallback_match.group(1).strip()
                # Stop bij "Project type"
                omschrijving = re.split(r'\s+Project\s+type', omschrijving, flags=re.IGNORECASE)[0].strip()
                if 3 <= len(omschrijving) <= 50 and not re.match(r'^\d', omschrijving):
                    result["omschrijving"] = omschrijving
                    break
    
    # ========== 4. REFERENTIENR. = KLANTREFERENTIE ==========
    # Zoek naar label "Referentienr." en haal waarde ernaast op (alleen als er echt iets staat)
    # Patroon: "Referentienr.: [waarde]" → klantreferentie = [waarde]
    # BELANGRIJK: Als het veld leeg is, moet dit None blijven (niet uitlezen)
    
    referentie_patterns = [
        r'Referentienr\.?[:\s]+([A-Za-z0-9]{3,})',  # Referentienr.: [waarde]
        r'Fase\s+referentie\s+nr\.?[:\s]+([A-Za-z0-9]{3,})',  # Fase referentie nr.: [waarde]
        r'Referentie[:\s]+([A-Za-z0-9]{3,})',  # Referentie: [waarde]
    ]
    
    # Alleen uitlezen als er een duidelijke waarde is (niet leeg)
    for pattern in referentie_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            ref_num = match.group(1).strip()
            # Filter lege waarden - als het alleen whitespace of speciale tekens is, skip
            if ref_num and len(ref_num) >= 3 and not re.match(r'^[\s\-_\.]+$', ref_num):
                result["referentienummer"] = ref_num
                break
    
    # Als geen specifieke referentie gevonden, blijf None (niet uitlezen)
    
    # Zoek datums - BELANGRIJK: onderscheid leverdatum en startdatum
    # Leverdatum staat vaak als "Fase opleverdatum" of "Opleverdatum" 
    # Startdatum staat vaak als "Start-/Opleverdatum" eerste deel of "Fase datum"
    
    # ========== 5. OPLEVERDATUM = LEVERDATUM ==========
    # Zoek naar label "Opleverdatum" of "Fase opleverdatum" en haal datum ernaast op
    # Patronen: "Opleverdatum: 19-12-2025", "Fase opleverdatum: 19-12-2025"
    
    opleverdatum_patterns = [
        r'Fase\s+opleverdatum[:\s]+(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # Fase opleverdatum: 19-12-2025
        r'Fase\s+opleverdatum[:\s]+(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})',  # Fase opleverdatum: 2025-12-19
        r'Opleverdatum[:\s]+(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # Opleverdatum: 19-12-2025
        r'Opleverdatum[:\s]+(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})',  # Opleverdatum: 2025-12-19
        r'Leverdatum[:\s]+(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # Leverdatum: 19-12-2025
        r'Leverdatum[:\s]+(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})',  # Leverdatum: 2025-12-19
    ]
    
    leverdatum_found = None
    for pattern in opleverdatum_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                if len(match.group(1)) == 4:  # yyyy-mm-dd formaat
                    jaar = match.group(1)
                    maand = match.group(2).zfill(2)
                    dag = match.group(3).zfill(2)
                    leverdatum_found = f"{jaar}-{maand}-{dag}"
                else:  # dd-mm-yyyy formaat (Nederlandse notatie: dag-maand-jaar!)
                    dag = match.group(1).zfill(2)
                    maand = match.group(2).zfill(2)
                    jaar = match.group(3)
                    # Validatie: dag moet 1-31 zijn, maand 1-12
                    if 1 <= int(dag) <= 31 and 1 <= int(maand) <= 12:
                        leverdatum_found = f"{jaar}-{maand}-{dag}"
                if leverdatum_found:
                    break
            except:
                pass
    
    # Zoek startdatum/fase datum - BELANGRIJK: "Fase datum" of eerste deel van "Start-/Opleverdatum"
    startdatum_patterns = [
        # "Fase datum" (specifiekste veld)
        r'Fase\s+datum[:\s]+(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # Fase datum: 14-11-2025
        r'Fase\s+datum[:\s]+(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})',  # Fase datum: 2025-11-14
        # "Start-/Opleverdatum" eerste deel (voor streepje)
        r'Start[-/]\s*Opleverdatum[:\s]+(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # Start-/Opleverdatum: 14-11-2025 (eerste deel)
        # Algemene startdatum
        r'(?:Startdatum|Start)[:\s]+(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # Startdatum: 14-11-2025
        r'(?:Startdatum|Start)[:\s]+(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})',  # Startdatum: 2025-11-14
    ]
    
    startdatum_found = None
    for pattern in startdatum_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                if len(match.group(1)) == 4:  # yyyy-mm-dd formaat
                    jaar = match.group(1)
                    maand = match.group(2).zfill(2)
                    dag = match.group(3).zfill(2)
                    startdatum_found = f"{jaar}-{maand}-{dag}"
                else:  # dd-mm-yyyy formaat (Nederlandse notatie: dag-maand-jaar!)
                    dag = match.group(1).zfill(2)
                    maand = match.group(2).zfill(2)
                    jaar = match.group(3)
                    # Validatie: dag moet 1-31 zijn, maand 1-12
                    if 1 <= int(dag) <= 31 and 1 <= int(maand) <= 12:
                        startdatum_found = f"{jaar}-{maand}-{dag}"
                if startdatum_found:
                    break
            except:
                pass
    
    # Als geen specifieke labels gevonden, zoek alle datums in tekst
    if not leverdatum_found or not startdatum_found:
        date_patterns = [
            r'(\d{1,2})[-\/](\d{1,2})[-\/](\d{4})',  # dd-mm-yyyy (Nederlandse notatie!)
            r'(\d{4})[-\/](\d{1,2})[-\/](\d{1,2})',  # yyyy-mm-dd
        ]
        
        dates_found = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                try:
                    if len(match[0]) == 4:  # yyyy-mm-dd
                        jaar = match[0]
                        maand = match[1].zfill(2)
                        dag = match[2].zfill(2)
                        dates_found.append(f"{jaar}-{maand}-{dag}")
                    else:  # dd-mm-yyyy (Nederlandse notatie: dag-maand-jaar!)
                        dag = match[0].zfill(2)
                        maand = match[1].zfill(2)
                        jaar = match[2]
                        # Validatie: dag moet 1-31 zijn, maand 1-12 (voorkomt omdraaien)
                        try:
                            dag_int = int(dag)
                            maand_int = int(maand)
                            if 1 <= dag_int <= 31 and 1 <= maand_int <= 12:
                                dates_found.append(f"{jaar}-{maand}-{dag}")
                        except:
                            pass
                except:
                    pass
        
        # Als we 2 datums vinden, is de laatste meestal leverdatum (later in tijd)
        if dates_found:
            if len(dates_found) >= 2:
                # Sorteer op datum - laatste (nieuwste) is meestal leverdatum
                sorted_dates = sorted(dates_found)
                # Als startdatum nog niet gevonden, gebruik eerste (oudste)
                if not startdatum_found:
                    startdatum_found = sorted_dates[0]
                # Als leverdatum nog niet gevonden, gebruik laatste (nieuwste)
                if not leverdatum_found:
                    leverdatum_found = sorted_dates[-1]
                # Als beide al gevonden, gebruik alleen als ze logisch zijn (leverdatum > startdatum)
                elif startdatum_found and leverdatum_found:
                    if leverdatum_found <= startdatum_found:
                        # Als leverdatum eerder is dan startdatum, wissel ze om
                        leverdatum_found, startdatum_found = sorted_dates[-1], sorted_dates[0]
            elif len(dates_found) == 1:
                # Alleen 1 datum gevonden - gebruik als leverdatum (meestal de belangrijkste)
                if not leverdatum_found:
                    leverdatum_found = dates_found[0]
    
    # Zet datums in resultaat
    if leverdatum_found:
        result["leverdatum"] = leverdatum_found
    if startdatum_found:
        result["startdatum"] = startdatum_found
    
    return result


# Basis extractie zonder OCR - probeer metadata te gebruiken
def extract_order_info_basic(image_data: bytes) -> dict:
    """Basis extractie zonder OCR - probeer wat te extraheren"""
    result = {
        "projectnummer": None,
        "omschrijving": None,
        "klant": None,
        "leverdatum": None,
        "startdatum": None,
        "referentienummer": None,
        "message": "OCR functionaliteit vereist pytesseract voor betere resultaten. Installeer met: pip install pytesseract"
    }
    
    # Probeer nog steeds wat basis informatie te extraheren van de filename of metadata
    # (Dit is een placeholder - in productie zou je hier betere extractie kunnen doen)
    
    return result

